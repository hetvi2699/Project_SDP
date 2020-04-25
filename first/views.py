import os
from django.shortcuts import render
from django.views.generic import TemplateView
from django.http import HttpResponseRedirect
from django.template.context_processors import csrf
from first.views import *
from first.templates import *
from first.models import *
from django.contrib.auth.models import User, Group
from django.contrib import auth
from django.contrib import messages
from django.contrib.auth import authenticate, login
from django.urls import path, include
from docx import Document
from docx.shared import Inches
import datetime
import pdfkit
import random
from django.shortcuts import render, redirect
from django.conf import settings
from django.contrib import messages
import requests
from win32com import client
from django.views.decorators.csrf import csrf_exempt
from first.templates.paytm import Checksum
from django.http import HttpResponse
from wsgiref.util import FileWrapper
from django.core.mail import send_mail
from django.template.loader import get_template
from django.template import Context
from django.core.mail import EmailMultiAlternatives
import json
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_protect


MERCHANT_KEY = 'UnWpilDBhAs2huGE'


def login(request):
    c = {}
    c.update(csrf(request))
    return render(request, 'login.html', c)


def new_registration1(request):
    c = {}
    c.update(csrf(request))
    return render(request, 'registration1.html', c)


def registration1(request):
    doc = Document()
    AppId = random.randint(000000000000, 999999999999)
    docx_title = "Reciept"
    doc.add_heading("RECEIPT FOR YOUR APPLICATION")
    doc.add_paragraph("")
    doc.add_paragraph("Application Id : " + '%d' % AppId)
    username = request.POST.get('name')
    doc.add_paragraph("Student Name: " + username)
    fname = request.POST.get('fname')
    mothername = request.POST.get('mname')
    dob = request.POST.get('dob')
    gender = request.POST.get('gender')
    rollno = request.POST.get('rollno')
    doc.add_paragraph("Roll No: " + rollno)
    category = request.POST.get('category')
    pwd = request.POST.get('pwd')
    password1 = request.POST.get('password1')
    password2 = request.POST.get('password2')
    if password1!=password2:
        return render(request, 'registration1.html', {"AppId": AppId})
    sque = request.POST.get('sque')
    sans = request.POST.get('sans')
    request.session['AppId'] = AppId
    s = Register(AppId=AppId, Name=username, FatherName=fname, MotherName=mothername, DOB=dob, Gender=gender,
                 Rollno=rollno, Category=category, PwdStatus=pwd, Password=password1, SecureQue=sque, SecureAns=sans)
    s.save()
    doc.save('Receipt.docx')
    return render(request, 'registration2.html', {"AppId": AppId})


def new_registration2(request):
    c = {}
    c.update(csrf(request))
    return render(request, 'registration2.html', c)


def registration2(request):
    res = request.POST.get('res')
    add1 = request.POST.get('add1')
    add2 = request.POST.get('add2')
    city = request.POST.get('city')
    district = request.POST.get('district')
    pincode = request.POST.get('pincode')
    state = request.POST.get('state')
    nationality = request.POST.get('nationality')
    email = request.POST.get('email')
    mob_no = request.POST.get('mob_no')
    AppId = request.session['AppId']
    s = Contact(AppId=AppId, Residence=res, Address1=add1, Address2=add2, City=city, District=district,
                Pincode=pincode, State=state, Nationality=nationality, EmailId=email, MobileNo=mob_no)
    s.save()
    return render(request, 'registration3.html', {"AppId": AppId})


def new_registration3(request):
    c = {}
    c.update(csrf(request))
    return render(request, 'registration3.html', c)


def registration3(request):
    tenth = 'Pass'
    tenth_year = request.POST.get('10_year')
    tenth_board = request.POST.get('10_board')
    tenth_roll = request.POST.get('10_roll')
    tenth_mode = request.POST.get('10_mode')
    tenth_res = request.POST.get('10_res')
    twelfth = 'Pass'
    twelfth_year = request.POST.get('12_year')
    twelfth_board = request.POST.get('12_board')
    twelfth_roll = request.POST.get('12_roll')
    twelfth_mode = request.POST.get('12_mode')
    twelfth_res = request.POST.get('12_res')
    AppId = request.session['AppId']
    s = Education(AppId=AppId, TenthStatus=tenth, TenthYear=tenth_year, TenthBoard=tenth_board, TenthRollNo=tenth_roll, TenthResultMode=tenth_mode, TenthResult=tenth_res,
                  TwelfthStatus=twelfth, TwelfthYear=twelfth_year, TwelfthBoard=twelfth_board, TwelfthRollNo=twelfth_roll, TwelfthResultMode=twelfth_mode, TwelfthResult=twelfth_res)
    s.save()
    return render(request, 'registration4.html', {"AppId": AppId})


def new_registration4(request):
    c = {}
    c.update(csrf(request))
    return render(request, 'registration4.html', c)


def registration4(request):
    mypic = request.FILES.get('mypic')
    mysign = request.FILES.get('mysign')
    print(mypic)
    print(mysign)
    AppId = request.session['AppId']
    s = Upload(AppId=AppId, Photo=mypic, Signature=mysign)
    s.save()
    f = open('Receipt.docx', 'rb')
    doc = Document(f)
    doc.add_paragraph("Payment  : 10 RS./- ")
    doc.add_paragraph("Exam Date : 24/04/2020 ")
    doc.add_paragraph("Exam Time : 5:30 pm ")
    f.close()
    doc.save('Receipt.docx')
    res = Contact.objects.filter(AppId=AppId).first()
    subject, from_email, to = 'Receipt for the examination', settings.EMAIL_HOST_USER, res.EmailId
    text_content = 'Keep it carefully'
    msg = EmailMultiAlternatives(subject, text_content, from_email, [to])
    msg.attach_file(
        'C:/Users/Manish Prajapati/Desktop/SDP_HETVI/firsttry/Receipt.docx')
    msg.send()
    return render(request, 'payment.html', {"AppId": AppId})


def convert_to_pdf(doc):
    f = open('test13.docx', 'rb')
    doc = Document(f)
    try:
        word = client.DispatchEx("Word.Application")
        word.Visible = True
        new_name = doc.replace(".docx", r".pdf")
        worddoc = word.Documents.Open(doc)
        print("Exporting ... ")
        worddoc.SaveAs(new_name, FileFormat=17)
        worddoc.Close()
        word.Quit()
    finally:
        print("hello")


def payment(request):
    c = {}
    c.update(csrf(request))
    AppId = request.session['AppId']
    request.session['AppId'] = AppId
    return render(request, 'payment.html', {"AppId": AppId})


def paymenterror(request):
    c = {}
    c.update(csrf(request))
    return render(request, 'paymenterror', {"AppId": AppId})


filename = 'Instructions.pdf'


def auth_view(request):
    if request.method == 'POST':
        username = request.POST['Username']
        password = request.POST['Password']
        print(username)
        student = Register.objects.filter(
            AppId=username, Password=password).exists()
        print(student)
        if student is not False:
            print(student)
            AppId = username
            res = Contact.objects.filter(AppId=AppId).first()
            subject, from_email, to = 'Introduction about the examination', settings.EMAIL_HOST_USER, res.EmailId
            text_content = 'Read the instructions carefully'
            msg = EmailMultiAlternatives(
                subject, text_content, from_email, [to])
            msg.attach_file(
                'C:/Users/Manish Prajapati/Desktop/SDP_HETVI/firsttry/media/Instructions.pdf')
            msg.send()
            request.session['AppId'] = AppId
        # return redirect('/first/certificate/')
            return render(request, 'welcome.html', {"AppId": AppId})
        else:
            Usernm="HetviPrajapati"
            Pwd="Hetvi123@#"
            #superusers = User.objects.filter(is_superuser=True)
            #superusers_pwds = User.objects.filter(is_superuser=True).values_list('password')
            #print(superusers_pwds)
            #print(superusers)
            print(username)
            if username==Usernm and Pwd==password:
                questions = Question.objects.all()
                return render(request, "show.html", {'questions': questions})
    return render(request, 'login.html')


def welcome(request):
    c = {}
    c.update(csrf(request))
    AppId = request.session['AppId']
    request.session['AppId'] = AppId
    return render(request, 'welcome.html', {"AppId": AppId})


def exam_start(request):
    AppId = request.session['AppId']
    request.session['AppId'] = AppId
    print(AppId)
    l = []
    for t in Question.objects.all():
        temp = {}
        # temp["SubjectId"] = t.SubjectId
        # temp["QuestionId"] = t.QuestionId
        temp["Question"] = t.Question
        temp["Option1"] = t.Option1
        temp["Option2"] = t.Option2
        temp["Option3"] = t.Option3
        temp["Option4"] = t.Option4
        temp["Answer"] = t.Answer
        l.append(temp)
        print(l)
        # context=serializers.serialize('json', l)
    return JsonResponse(l, safe=False)


def examreq(request):
    # c = {}
    # c.update(csrf(request))
    AppId = request.session['AppId']
    request.session['AppId'] = AppId
    return render(request, 'exam.html', {"AppId": AppId})


def examcompleted(request):
    c = {}
    c.update(csrf(request))
    AppId = request.session['AppId']
    request.session['AppId'] = AppId
    return redirect(request, 'examcompleted.html', {"AppId": AppId})


def certi(request):
    c = {}
    c.update(csrf(request))
    return render(request, 'certi.html')


def exampro(request):
    c = {}
    c.update(csrf(request))
    que_id = request.GET.get("id", False)
    temp = Question.objects.get(QuestionId=que_id)
    print(temp)
    return render(request, 'exampro.html')


def download_pdf(request):
    projectUrl = request.get_host() + '/first/certificate'
    config = pdfkit.configuration(wkhtmltopdf='/opt/bin/wkhtmltopdf')
    pdf = pdfkit.from_url(projectUrl, False, configuration=config)
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="ourcodeworld.pdf"'
    #pdfkit.from_string("<h1>Hello World</h1>", ourcodeworld.pdf, configuration=config)
    return HttpResponse(response)


def count(request):
    AppId = request.session['AppId']
    if request.method == 'POST':
        res = Result.objects.get(AppId=AppId)
        #res.CorrectAns =  request.POST.get()
        #res.IncorrectAns = request.POST['incorrectAnswers']
        res.Result = request.POST.get('data')
        res.save()
        print(res)
    return response(request, 'new_certificate', {"AppId": AppId})


def new_certificate(request):
    c = {}
    c.update(csrf(request))
    AppId = request.session['AppId']
    request.session['AppId'] = AppId
    return redirect(request, 'certificate.html', {"AppId": AppId})


def certificate(request):
    # aid=request.GET.get('AppId')
    AppId = request.session['AppId']
    print(AppId)
    if request.method == 'POST':
        res = Result.objects.filter(AppId=AppId).first()
        print(res)
        new = request.POST.get('data')
        print(new)
        res.AppId = AppId
        res.CorrectAns = new.correctAnswers
        res.IncorrectAns = new.incorrectAnswers
        res.Result = new.cs
        res.Rank = 100
        res.save()
        print(res)
    student_temp = Register.objects.get(AppId=AppId)
    print(student_temp)
    result_temp = Result.objects.get(AppId=AppId)
    print(result_temp)
    context = {'AppId': AppId, 'student': student_temp, 'result': result_temp}
    return render(request, 'certificate.html', context)


def addque(request):
    c = {}
    c.update(csrf(request))
    return render(request, "addque.html")


def enterque(request):
    qid = request.POST.get('QuestionId')
    que = request.POST.get('Question')
    op1 = request.POST.get('Option1')
    op2 = request.POST.get('Option2')
    op3 = request.POST.get('Option3')
    op4 = request.POST.get('Option4')
    ans = request.POST.get('Answer')
    p = Question(
        QuestionId=qid,
        Question=que,
        Option1=op1,
        Option2=op2,
        Option3=op3,
        Option4=op4,
        Answer=ans,
    )
    p.save()
    questions = Question.objects.all()
    return render(request, "show.html", {'questions': questions})


def show(request):
    questions = Question.objects.all()
    return render(request, "show.html", {'questions': questions})


def edit(request, id):
    question = Question.objects.get(QuestionId=id)
    return render(request, 'edit.html', {'question': question})


def update(request, id):
    #question = Question.objects.get(QuestionId=id)
    qid = request.POST.get('queid')
    que = request.POST.get('que')
    op1 = request.POST.get('op1')
    op2 = request.POST.get('op2')
    op3 = request.POST.get('op3')
    op4 = request.POST.get('op4')
    ans = request.POST.get('ans')

    dbEntry = Question.objects.filter(QuestionId=id).first()
    if dbEntry:
        dbEntry.QuestionId = qid
        dbEntry.Question = que
        dbEntry.Option1 = op1
        dbEntry.Option2 = op2
        dbEntry.Option3 = op3
        dbEntry.Option4 = op4
        dbEntry.Answer = ans
        dbEntry.save()
    questions = Question.objects.all()
    return render(request, "show.html", {'questions': questions})


def destroy(request, id):
    question = Question.objects.get(QuestionId=id)
    question.delete()
    questions = Question.objects.all()
    return render(request, "show.html", {'questions': questions})


def logout(request):
    c = {}
    c.update(csrf(request))
    return render(request, 'login.html')


def frgtpwd(request):
    AppId = request.POST.get('AppId')
    pwd = request.POST.get('password')
    pwd1 = request.POST.get('password1')
    if pwd1!=pwd:
        return render(request, 'frgtpwd.html', {"AppId": AppId})
    sque = request.POST.get('sque')
    sans = request.POST.get('sans')
    student = Register.objects.filter(
        AppId=AppId, SecureQue=sque, SecureAns=sans).first()
    if student:
        print(student)
        student.Password = pwd
        student.save()
        res = Contact.objects.filter(AppId=AppId).first()
        subject, from_email, to = 'Forgot Password ', settings.EMAIL_HOST_USER, res.EmailId
        text_content = 'Your Password For the online apptitude examination changed sucessfully'
        send_mail(subject, text_content, from_email, [to], fail_silently=False)
        request.session['AppId'] = AppId
        return render(request, 'login.html', {"AppId": AppId})
    else:
        #res = Contact.objects.filter(AppId=AppId).first()
        #print(res)
        #subject, from_email, to = 'Password Change Activity Suspected ', settings.EMAIL_HOST_USER, res.EmailId
        #text_content = 'Any anonymous user tries to change your password for the online apptitude examination. You should check this activity'
        #send_mail(subject, text_content, from_email, [to], fail_silently=False)
        return render(request, 'frgtpwd.html')

def contactus(request):
    AppId = request.POST.get('AppId')
    email = request.POST.get('email')
    subject = request.POST.get('subject')
    message = request.POST.get('message')
    from_email = settings.EMAIL_HOST_USER
    send_mail(subject, message, email, [from_email], fail_silently=False)
    return render(request, 'welcome.html')
        
